import streamlit as st
from docx import Document
import spacy
from transformers import BertTokenizer, BertForSequenceClassification
import torch

# Load pre-trained NLP model for resume-job matching (e.g., BERT or Sentence-BERT)
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertForSequenceClassification.from_pretrained('bert-base-uncased')

# Function to extract text from a Word document
def extract_text_from_docx(file):
    doc = Document(file)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

# Function to compare resume with job description using BERT
def compare_resume_to_job(resume_text, job_description):
    inputs = tokenizer.encode_plus(resume_text, job_description, return_tensors='pt', padding=True, truncation=True, max_length=512)
    outputs = model(**inputs)
    logits = outputs.logits
    prediction = torch.argmax(logits, dim=-1).item()
    return prediction

# Streamlit UI
st.title("LLM-Based Resume Checker")

# Upload resume
uploaded_file = st.file_uploader("Upload Your Resume (Word Doc)", type=["docx"])
if uploaded_file:
    resume_text = extract_text_from_docx(uploaded_file)
    st.write("Resume content extracted: ", resume_text[:500])  # Show a preview of the resume content

    # Job description input
    job_desc = st.text_area("Paste Job Description Here", height=150)

    if job_desc:
        # Compare resume to job description
        result = compare_resume_to_job(resume_text, job_desc)
        
        if result == 1:  # Assuming 1 is a match and 0 is not
            st.success("The resume matches the job description well!")
        else:
            st.warning("The resume does not match the job description perfectly. Suggestions:")
            # Placeholder for suggestions (you could integrate NLP for detailed feedback)
            st.write("Consider adding these missing skills: 'Python', 'Machine Learning', etc.")

        # Ask if the user wants the resume modified
        modify_button = st.button("Do you want me to modify the resume?")
        if modify_button:
            st.write("Modifying the resume...")
            # Here you would implement resume modification logic
            modified_resume_text = resume_text + "\n\nAdditional skills: Python, Machine Learning"
            st.write("Modified resume: ", modified_resume_text[:500])

            # Generate modified Word document
            doc = Document()
            doc.add_paragraph(modified_resume_text)
            modified_filename = "modified_resume.docx"
            doc.save(modified_filename)

            # Provide download link
            st.download_button("Download Modified Resume", modified_filename)

